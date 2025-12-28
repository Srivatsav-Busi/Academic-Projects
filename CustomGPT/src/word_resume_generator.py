"""
Word Document Generator for Resumes.
Creates professional Word (.docx) resumes from markdown content.
"""

import io
import logging
from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

logger = logging.getLogger(__name__)

class WordResumeGenerator:
    """Generator for creating Word document resumes."""
    
    def __init__(self):
        self.doc = None
        
    def create_resume_document(self, resume_data: Dict[str, Any]) -> Optional[bytes]:
        """Create a Word document from resume data."""
        if not WORD_AVAILABLE:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return None
            
        try:
            # Create new document
            self.doc = Document()
            
            # Set up styles
            self._setup_styles()
            
            # Add header
            self._add_header(resume_data)
            
            # Add sections
            self._add_summary(resume_data.get('summary', ''))
            self._add_technical_skills(resume_data.get('technical_skills', []))
            self._add_experience(resume_data.get('experience', []))
            self._add_education(resume_data.get('education', []))
            self._add_projects(resume_data.get('projects', []))
            self._add_certifications(resume_data.get('certifications', []))
            
            # Save to bytes
            doc_buffer = io.BytesIO()
            self.doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to create Word document: {e}")
            return None
    
    def _setup_styles(self):
        """Set up document styles."""
        styles = self.doc.styles
        
        # Header style
        if 'Header' not in [style.name for style in styles]:
            header_style = styles.add_style('Header', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Arial'
            header_font.size = Pt(24)
            header_font.bold = True
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Section style
        if 'Section' not in [style.name for style in styles]:
            section_style = styles.add_style('Section', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Arial'
            section_font.size = Pt(14)
            section_font.bold = True
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
        
        # Body style
        if 'Body' not in [style.name for style in styles]:
            body_style = styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Arial'
            body_font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(6)
    
    def _add_header(self, resume_data: Dict[str, Any]):
        """Add resume header with name and contact info."""
        # Name
        name_para = self.doc.add_paragraph()
        name_para.style = 'Header'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(resume_data.get('name', 'Srivatsav Busi'))
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        
        # Contact info
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.style = 'Body'
        
        contact_info = [
            resume_data.get('phone', '4708909738'),
            resume_data.get('email', 'srivatsavbusi@gmail.com'),
            resume_data.get('location', 'Atlanta, GA'),
            resume_data.get('linkedin', 'https://linkedin.com/in/srivatsavbusi'),
            resume_data.get('github', 'https://github.com/srivatsavbusi')
        ]
        
        contact_text = ' | '.join([info for info in contact_info if info])
        contact_para.add_run(contact_text)
        
        # Add line break
        self.doc.add_paragraph()
    
    def _add_summary(self, summary: str):
        """Add professional summary section."""
        if not summary:
            return
            
        self.doc.add_paragraph('PROFESSIONAL SUMMARY', style='Section')
        
        summary_para = self.doc.add_paragraph()
        summary_para.style = 'Body'
        summary_para.add_run(summary)
    
    def _add_technical_skills(self, skills: list):
        """Add technical skills section."""
        if not skills:
            return
            
        self.doc.add_paragraph('TECHNICAL SKILLS', style='Section')
        
        # Group skills by category
        skill_categories = {
            'Programming Languages': [],
            'Frameworks & Libraries': [],
            'Tools & Technologies': [],
            'Cloud Platforms': [],
            'Databases': []
        }
        
        for skill in skills:
            skill_lower = skill.lower()
            if any(lang in skill_lower for lang in ['python', 'r', 'sql', 'pyspark', 'shell', 'bash']):
                skill_categories['Programming Languages'].append(skill)
            elif any(fw in skill_lower for fw in ['spark', 'pandas', 'numpy', 'flask', 'scikit', 'tensorflow', 'pytorch']):
                skill_categories['Frameworks & Libraries'].append(skill)
            elif any(tool in skill_lower for tool in ['airflow', 'nifi', 'tableau', 'git', 'terraform', 'docker', 'kubernetes']):
                skill_categories['Tools & Technologies'].append(skill)
            elif any(cloud in skill_lower for cloud in ['aws', 'gcp', 'azure', 'bigquery', 'vertex', 'emr']):
                skill_categories['Cloud Platforms'].append(skill)
            elif any(db in skill_lower for db in ['mongodb', 'cassandra', 'redshift', 'dynamodb', 'mysql', 'postgresql']):
                skill_categories['Databases'].append(skill)
            else:
                skill_categories['Tools & Technologies'].append(skill)
        
        for category, category_skills in skill_categories.items():
            if category_skills:
                skills_para = self.doc.add_paragraph()
                skills_para.style = 'Body'
                skills_para.add_run(f"• {category}: ").bold = True
                skills_para.add_run(', '.join(category_skills))
    
    def _add_experience(self, experience: list):
        """Add work experience section."""
        if not experience:
            return
            
        self.doc.add_paragraph('EXPERIENCE', style='Section')
        
        for exp in experience:
            # Company and position
            exp_para = self.doc.add_paragraph()
            exp_para.style = 'Body'
            exp_para.add_run(f"{exp.get('position', '')}").bold = True
            exp_para.add_run(f" | {exp.get('company', '')} | {exp.get('location', '')} | {exp.get('duration', '')}")
            
            # Responsibilities
            if exp.get('responsibilities'):
                for responsibility in exp['responsibilities']:
                    resp_para = self.doc.add_paragraph()
                    resp_para.style = 'Body'
                    resp_para.add_run(f"• {responsibility}")
    
    def _add_education(self, education: list):
        """Add education section."""
        if not education:
            return
            
        self.doc.add_paragraph('EDUCATION', style='Section')
        
        for edu in education:
            edu_para = self.doc.add_paragraph()
            edu_para.style = 'Body'
            edu_para.add_run(f"{edu.get('degree', '')}").bold = True
            edu_para.add_run(f" | {edu.get('institution', '')} | {edu.get('location', '')} | {edu.get('year', '')}")
    
    def _add_projects(self, projects: list):
        """Add projects section."""
        if not projects:
            return
            
        self.doc.add_paragraph('PROJECTS', style='Section')
        
        for project in projects:
            proj_para = self.doc.add_paragraph()
            proj_para.style = 'Body'
            proj_para.add_run(f"{project.get('name', '')}").bold = True
            proj_para.add_run(f" | {project.get('description', '')}")
    
    def _add_certifications(self, certifications: list):
        """Add certifications section."""
        if not certifications:
            return
            
        self.doc.add_paragraph('CERTIFICATIONS', style='Section')
        
        for cert in certifications:
            cert_para = self.doc.add_paragraph()
            cert_para.style = 'Body'
            cert_para.add_run(f"• {cert}")
    
    def create_resume_from_markdown(self, markdown_content: str, job_title: str, company: str) -> Optional[bytes]:
        """Create a Word resume from markdown content."""
        try:
            # Parse markdown content to extract resume data
            resume_data = self._parse_markdown_resume(markdown_content, job_title, company)
            
            # Create Word document
            return self.create_resume_document(resume_data)
            
        except Exception as e:
            logger.error(f"Failed to create resume from markdown: {e}")
            return None
    
    def _parse_markdown_resume(self, markdown_content: str, job_title: str, company: str) -> Dict[str, Any]:
        """Parse markdown resume content to extract structured data."""
        lines = markdown_content.split('\n')
        
        resume_data = {
            'name': 'Srivatsav Busi',
            'phone': '4708909738',
            'email': 'srivatsavbusi@gmail.com',
            'location': 'Atlanta, GA',
            'linkedin': 'https://linkedin.com/in/srivatsavbusi',
            'github': 'https://github.com/srivatsavbusi',
            'summary': '',
            'technical_skills': [],
            'experience': [],
            'education': [],
            'projects': [],
            'certifications': []
        }
        
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Parse sections
            if line.startswith('##'):
                current_section = line.replace('##', '').strip().lower()
            elif line.startswith('#'):
                current_section = line.replace('#', '').strip().lower()
            elif line.startswith('**') and line.endswith('**'):
                # Bold text - likely section headers
                current_section = line.replace('**', '').strip().lower()
            elif current_section == 'summary' or 'summary' in current_section:
                if line.startswith('•') or line.startswith('-'):
                    resume_data['summary'] += line[1:].strip() + ' '
                else:
                    resume_data['summary'] += line + ' '
            elif 'skill' in current_section:
                if line.startswith('•') or line.startswith('-'):
                    skill = line[1:].strip()
                    if skill:
                        resume_data['technical_skills'].append(skill)
            elif 'experience' in current_section:
                if line.startswith('•') or line.startswith('-'):
                    # This is a responsibility
                    if resume_data['experience']:
                        resume_data['experience'][-1]['responsibilities'].append(line[1:].strip())
            elif 'education' in current_section:
                if line and not line.startswith('#'):
                    # This is education entry
                    resume_data['education'].append({
                        'degree': line,
                        'institution': '',
                        'location': '',
                        'year': ''
                    })
            elif 'certification' in current_section:
                if line.startswith('•') or line.startswith('-'):
                    resume_data['certifications'].append(line[1:].strip())
        
        # Clean up summary
        resume_data['summary'] = resume_data['summary'].strip()
        
        return resume_data

def create_word_resume(markdown_content: str, job_title: str, company: str) -> Optional[bytes]:
    """Create a Word resume from markdown content."""
    generator = WordResumeGenerator()
    return generator.create_resume_from_markdown(markdown_content, job_title, company)







